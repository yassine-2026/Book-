import os, uuid, json, io, zipfile, time
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from groq import Groq
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageColor
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

OUTPUT_DIR = 'static/books'
COVERS_DIR = 'static/covers'
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(COVERS_DIR, exist_ok=True)

# ========== 30+ LANGUAGES ==========
LANGUAGES = {
    'arabic': 'العربية', 'english': 'English', 'french': 'Français',
    'spanish': 'Español', 'german': 'Deutsch', 'italian': 'Italiano',
    'portuguese': 'Português', 'russian': 'Русский', 'chinese': '中文',
    'japanese': '日本語', 'korean': '한국어', 'hindi': 'हिन्दी',
    'turkish': 'Türkçe', 'urdu': 'اردو', 'malay': 'Bahasa Melayu',
    'indonesian': 'Bahasa Indonesia', 'dutch': 'Nederlands',
    'polish': 'Polski', 'swedish': 'Svenska', 'norwegian': 'Norsk',
    'danish': 'Dansk', 'finnish': 'Suomi', 'greek': 'Ελληνικά',
    'thai': 'ไทย', 'vietnamese': 'Tiếng Việt', 'hebrew': 'עברית',
    'persian': 'فارسی', 'bengali': 'বাংলা', 'tamil': 'தமிழ்',
    'tagalog': 'Tagalog', 'romanian': 'Română', 'ukrainian': 'Українська',
}

# ========== 15 COVER THEMES ==========
COVER_THEMES = {
    'educational': {'bg': '#1a237e', 'accent': '#ffd600', 'secondary': '#3949ab', 'pattern': 'books'},
    'story': {'bg': '#4a148c', 'accent': '#ff4081', 'secondary': '#7b1fa2', 'pattern': 'stars'},
    'self_help': {'bg': '#004d40', 'accent': '#00e5ff', 'secondary': '#00695c', 'pattern': 'sunrise'},
    'technology': {'bg': '#0d47a1', 'accent': '#00e676', 'secondary': '#1565c0', 'pattern': 'circuit'},
    'religious': {'bg': '#1b5e20', 'accent': '#ffd700', 'secondary': '#2e7d32', 'pattern': 'mosque'},
    'history': {'bg': '#3e2723', 'accent': '#ffab00', 'secondary': '#4e342e', 'pattern': 'scroll'},
    'cooking': {'bg': '#bf360c', 'accent': '#ffeb3b', 'secondary': '#d84315', 'pattern': 'food'},
    'business': {'bg': '#01579b', 'accent': '#ff6d00', 'secondary': '#0277bd', 'pattern': 'chart'},
    'science': {'bg': '#311b92', 'accent': '#00e5ff', 'secondary': '#4527a0', 'pattern': 'atoms'},
    'romance': {'bg': '#880e4f', 'accent': '#ff80ab', 'secondary': '#ad1457', 'pattern': 'hearts'},
    'horror': {'bg': '#212121', 'accent': '#ff1744', 'secondary': '#424242', 'pattern': 'dark'},
    'fantasy': {'bg': '#1a237e', 'accent': '#ffd740', 'secondary': '#283593', 'pattern': 'magic'},
    'comedy': {'bg': '#f57f17', 'accent': '#ffffff', 'secondary': '#f9a825', 'pattern': 'fun'},
    'travel': {'bg': '#006064', 'accent': '#18ffff', 'secondary': '#00838f', 'pattern': 'world'},
    'health': {'bg': '#1b5e20', 'accent': '#69f0ae', 'secondary': '#2e7d32', 'pattern': 'leaf'},
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/health')
def health_check():
    return jsonify({"status": "ok"}), 200

# ========== GENERATE BOOK ==========
@app.route('/api/generate-book', methods=['POST'])
def generate_book():
    try:
        data = request.get_json(force=True)
        topic = data.get('topic', '').strip()
        genre = data.get('genre', 'educational')
        chapters_count = min(int(data.get('chapters', 5)), 15)
        language = data.get('language', 'arabic')
        author = data.get('author', 'مؤلف').strip() or 'مؤلف'
        max_pages = min(int(data.get('max_pages', 100)), 100)
        
        if not topic:
            return jsonify({"success": False, "error": "Topic required"}), 400
        
        if not groq_client:
            return jsonify({"success": False, "error": "GROQ_API_KEY not configured"}), 500
        
        lang_name = LANGUAGES.get(language, 'English')
        theme = COVER_THEMES.get(genre, COVER_THEMES['educational'])
        
        # Step 1: Generate book structure
        prompt = f"""Create a COMPLETE professional {genre} book in {lang_name} about "{topic}".
Author: {author}
Chapters: {chapters_count}
Maximum pages: {max_pages}

CRITICAL RULES:
1. Each chapter MUST be 400-800 words of UNIQUE content
2. NO repetition between chapters
3. REAL, accurate information about the topic
4. Professional writing style
5. NO spelling or grammar errors
6. Each chapter should have practical examples

Return ONLY valid JSON (no other text):
{{
    "title": "Creative Book Title",
    "subtitle": "Engaging Subtitle",
    "description": "Compelling 2-3 sentence book description",
    "keywords": "keyword1, keyword2, keyword3",
    "chapters": [
        {{
            "number": 1,
            "title": "Chapter Title",
            "content": "Full chapter content 400-800 words with examples..."
        }}
    ]
}}
"""
        
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.8,
            max_tokens=8000
        )
        
        text = response.choices[0].message.content
        
        # Parse JSON
        start = text.find('{')
        end = text.rfind('}') + 1
        if start >= 0 and end > start:
            book_data = json.loads(text[start:end])
        else:
            return jsonify({"success": False, "error": "Failed to generate book"}), 500
        
        # Ensure minimum chapters
        if not book_data.get('chapters'):
            book_data['chapters'] = [{"number": 1, "title": "Introduction", "content": "Content here..."}]
        
        book_data['chapters'] = book_data['chapters'][:chapters_count]
        
        # Step 2: Generate professional cover
        book_id = uuid.uuid4().hex[:12]
        cover_path = create_cover(book_data['title'], book_data.get('subtitle', ''), author, genre, book_id)
        
        # Step 3: Create all formats
        os.makedirs(f'{OUTPUT_DIR}/{book_id}', exist_ok=True)
        
        pdf_path = create_pdf(book_data, author, book_id)
        docx_path = create_docx(book_data, author, book_id)
        txt_path = create_txt(book_data, author, book_id)
        html_path = create_html(book_data, author, book_id)
        
        # Step 4: Create ZIP
        zip_path = f'{OUTPUT_DIR}/{book_id}/all_formats.zip'
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(pdf_path, f'{book_data["title"]}.pdf')
            zf.write(docx_path, f'{book_data["title"]}.docx')
            zf.write(txt_path, f'{book_data["title"]}.txt')
            zf.write(html_path, f'{book_data["title"]}.html')
        
        return jsonify({
            "success": True,
            "book_id": book_id,
            "title": book_data['title'],
            "subtitle": book_data.get('subtitle', ''),
            "description": book_data.get('description', ''),
            "chapters": book_data['chapters'],
            "cover_url": f"/{cover_path}",
            "downloads": {
                "pdf": f"/{pdf_path}",
                "docx": f"/{docx_path}",
                "txt": f"/{txt_path}",
                "html": f"/{html_path}",
                "zip": f"/{zip_path}"
            }
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# ========== CREATE COVER ==========
def create_cover(title, subtitle, author, genre, book_id):
    theme = COVER_THEMES.get(genre, COVER_THEMES['educational'])
    
    # Create 1600x2400 cover (professional book ratio)
    img = Image.new('RGB', (1600, 2400), theme['bg'])
    draw = ImageDraw.Draw(img)
    
    # Background pattern
    if theme['pattern'] == 'books':
        for i in range(0, 1600, 80):
            draw.rectangle([i, 0, i+40, 2400], fill=theme['secondary'], outline=None)
    elif theme['pattern'] == 'stars':
        import random
        for _ in range(200):
            x, y = random.randint(0, 1600), random.randint(0, 2400)
            r = random.randint(1, 3)
            draw.ellipse([x-r, y-r, x+r, y+r], fill=theme['accent'])
    elif theme['pattern'] == 'circuit':
        for i in range(0, 1600, 60):
            draw.line([(i, 0), (i+30, 2400)], fill=theme['secondary'], width=1)
            for j in range(0, 2400, 60):
                draw.ellipse([i+25, j+25, i+35, j+35], fill=theme['accent'])
    
    # Gradient overlay
    for y in range(1200):
        alpha = int(180 * (1 - y/1200))
        try:
            color = ImageColor.getrgb(theme['accent'])
        except:
            color = (255, 255, 255)
        draw.line([(0, y), (1600, y)], fill=(color[0], color[1], color[2], alpha))
    
    # Decorative border
    draw.rectangle([40, 40, 1560, 2360], outline=theme['accent'], width=4)
    draw.rectangle([55, 55, 1545, 2345], outline=theme['accent'], width=1)
    
    # Title
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 72)
        font_sub = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 36)
        font_author = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 28)
    except:
        font_title = ImageFont.load_default()
        font_sub = font_title
        font_author = font_title
    
    # Draw title word by word (centered)
    words = title.split()
    y_pos = 300
    for word in words:
        bbox = draw.textbbox((0, 0), word, font=font_title)
        text_width = bbox[2] - bbox[0]
        x_pos = (1600 - text_width) / 2
        draw.text((x_pos, y_pos), word, fill=theme['accent'], font=font_title)
        y_pos += 80
    
    # Subtitle
    if subtitle:
        y_pos += 60
        bbox = draw.textbbox((0, 0), subtitle, font=font_sub)
        text_width = bbox[2] - bbox[0]
        draw.text(((1600-text_width)/2, y_pos), subtitle, fill='#ffffff', font=font_sub)
    
    # Author
    author_text = f'By: {author}'
    bbox = draw.textbbox((0, 0), author_text, font=font_author)
    text_width = bbox[2] - bbox[0]
    draw.text(((1600-text_width)/2, 2100), author_text, fill='#ffffff', font=font_author)
    
    # Genre badge
    badge_text = genre.upper()
    bbox = draw.textbbox((0, 0), badge_text, font=font_author)
    text_width = bbox[2] - bbox[0]
    draw.rectangle([(1600-text_width)/2-20, 2200, (1600+text_width)/2+20, 2250], fill=theme['accent'])
    draw.text(((1600-text_width)/2, 2205), badge_text, fill=theme['bg'], font=font_author)
    
    path = f'{COVERS_DIR}/cover_{book_id}.png'
    img.save(path, 'PNG')
    return path

# ========== CREATE PDF ==========
def create_pdf(book_data, author, book_id):
    path = f'{OUTPUT_DIR}/{book_id}/book.pdf'
    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=28, spaceAfter=30, alignment=TA_CENTER, textColor='#1a237e')
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontSize=18, spaceAfter=20, textColor='#4a148c')
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=12, spaceAfter=12, leading=18)
    
    story = []
    
    # Cover page
    story.append(Spacer(1, 200))
    story.append(Paragraph(book_data['title'], title_style))
    if book_data.get('subtitle'):
        story.append(Paragraph(book_data['subtitle'], ParagraphStyle('Sub', parent=title_style, fontSize=16, textColor='#666666')))
    story.append(Spacer(1, 50))
    story.append(Paragraph(f'By: {author}', ParagraphStyle('Author', alignment=TA_CENTER, fontSize=14)))
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph('Table of Contents', heading_style))
    story.append(Spacer(1, 20))
    for ch in book_data['chapters']:
        story.append(Paragraph(f"Chapter {ch['number']}: {ch['title']}", body_style))
    story.append(PageBreak())
    
    # Chapters
    page_count = 3
    for ch in book_data['chapters']:
        if page_count >= 100:
            break
        
        story.append(Paragraph(f"Chapter {ch['number']}", heading_style))
        story.append(Paragraph(ch['title'], ParagraphStyle('ChTitle', parent=heading_style, fontSize=14)))
        story.append(Spacer(1, 20))
        
        paragraphs = ch.get('content', '').split('\n')
        for p in paragraphs:
            if p.strip():
                story.append(Paragraph(p.strip(), body_style))
                page_count += 0.2
        
        story.append(PageBreak())
        page_count += 1
    
    doc.build(story)
    return path

# ========== CREATE DOCX ==========
def create_docx(book_data, author, book_id):
    path = f'{OUTPUT_DIR}/{book_id}/book.docx'
    doc = Document()
    
    # Title page
    doc.add_heading(book_data['title'], 0)
    if book_data.get('subtitle'):
        doc.add_heading(book_data['subtitle'], 1)
    doc.add_paragraph(f'By: {author}')
    doc.add_page_break()
    
    # TOC
    doc.add_heading('Table of Contents', 1)
    for ch in book_data['chapters']:
        doc.add_paragraph(f"Chapter {ch['number']}: {ch['title']}")
    doc.add_page_break()
    
    # Chapters
    for ch in book_data['chapters']:
        doc.add_heading(f"Chapter {ch['number']}: {ch['title']}", 1)
        doc.add_paragraph(ch.get('content', ''))
        doc.add_page_break()
    
    doc.save(path)
    return path

# ========== CREATE TXT ==========
def create_txt(book_data, author, book_id):
    path = f'{OUTPUT_DIR}/{book_id}/book.txt'
    with open(path, 'w', encoding='utf-8') as f:
        f.write(f"{book_data['title']}\n")
        f.write(f"By: {author}\n")
        f.write("="*50 + "\n\n")
        for ch in book_data['chapters']:
            f.write(f"CHAPTER {ch['number']}: {ch['title']}\n")
            f.write("-"*30 + "\n")
            f.write(ch.get('content', '') + "\n\n")
    return path

# ========== CREATE HTML ==========
def create_html(book_data, author, book_id):
    path = f'{OUTPUT_DIR}/{book_id}/book.html'
    with open(path, 'w', encoding='utf-8') as f:
        f.write(f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>{book_data['title']}</title>
<style>
body{{font-family:Arial,sans-serif;max-width:800px;margin:0 auto;padding:20px;line-height:1.8;background:#fafafa}}
h1{{color:#1a237e;text-align:center;font-size:2em}}
h2{{color:#4a148c;margin-top:40px}}
.chapter{{margin:30px 0;padding:25px;background:#fff;border-radius:10px;box-shadow:0 2px 10px rgba(0,0,0,0.1)}}
.cover{{text-align:center;padding:100px 0;background:linear-gradient(135deg,#1a237e,#4a148c);color:#fff;border-radius:20px;margin-bottom:50px}}
</style></head>
<body>
<div class="cover">
<h1 style="color:#fff">{book_data['title']}</h1>
<p style="font-size:1.2em">By: {author}</p>
</div>
""")
        for ch in book_data['chapters']:
            f.write(f'<div class="chapter"><h2>Chapter {ch["number"]}: {ch["title"]}</h2><p>{ch.get("content","")}</p></div>')
        f.write('</body></html>')
    return path

@app.route('/static/books/<path:filename>')
def serve_book(filename):
    return send_file(os.path.join(OUTPUT_DIR, filename))

@app.route('/static/covers/<path:filename>')
def serve_cover(filename):
    return send_file(os.path.join(COVERS_DIR, filename))

@app.route('/api/health')
def health():
    return jsonify({
        "status": "ok",
        "groq": bool(GROQ_API_KEY),
        "languages": len(LANGUAGES),
        "themes": len(COVER_THEMES)
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
